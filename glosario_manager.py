class GlossaryManager:
    def __init__(self, glossary=None):
        """Inicializa el GlossaryManager con un glosario opcional."""
        self.glosario = glossary if glossary is not None else {}

    def load_glossary(self, glossaryxml: str) -> dict:
        """Carga un glosario desde un archivo XML."""
        self.glosario = self.from_xml(glossaryxml)
        return self.glosario

    def get_term(self, term: str) -> dict:
        """Obtiene la definición y ejemplos de un término del glosario."""
        return self.glosario.get(term, {})

    def save_glossary(self, path: str):
        """Guarda el glosario actual en un archivo XML en la ruta especificada."""
        with open(path, 'w', encoding='utf-8') as f:
            f.write(self.to_xml("Autor", "Título"))

    def save_term(self, term: str, definition: str, examples: list):
        """Guarda un término en el glosario."""
        self.glosario[term] = {"definicion": definition, "ejemplos": examples}

    def delete_term(self, term: str):
        """Elimina un término del glosario."""
        if term in self.glosario:
            del self.glosario[term]

    def to_latex(self, author: str, title: str) -> str:
        """Convierte el glosario a un documento LaTeX."""
        latex_document = "\\documentclass[twocolumn]{article}\n"
        latex_document += "\\usepackage[utf8]{inputenc}\n"
        latex_document += "\\usepackage{enumitem}\n"
        latex_document += f"\\title{{{title}}}\n"
        latex_document += f"\\author{{{author}}}\n"
        latex_document += "\\date{}\n"
        latex_document += "\\begin{document}\n"
        latex_document += "\\maketitle\n"

        for term, data in self.glosario.items():
            latex_document += f"\\section*{{{term}}}\n"
            latex_document += f"\\textbf{{{data['definicion']}}}\n"
            latex_document += "\\begin{itemize}[label=\\textbullet]\n"
            for example in data["ejemplos"]:
                latex_document += f"\\item {example}\n"
            latex_document += "\\end{itemize}\n"

        latex_document += "\\end{document}\n"

        return latex_document

    def to_md(self, author: str, title: str) -> str:
        """Convierte el glosario a un documento Markdown."""
        md_document = f"# {title}\n\n"
        md_document += f"Autor: {author}\n\n"

        for term, data in self.glosario.items():
            md_document += f"## {term}\n\n"
            md_document += f"**Definición:** {data['definicion']}\n\n"
            md_document += "**Ejemplos:**\n\n"
            for example in data["ejemplos"]:
                md_document += f"- {example}\n"
            md_document += "\n"

        return md_document

    def to_html(self, author: str, title: str) -> str:
        """Convierte el glosario a un documento HTML."""
        html_document = "<!DOCTYPE html>\n<html>\n<head>\n"
        html_document += "<link rel='stylesheet' type='text/css' href='styles.css'>\n"
        html_document += "</head>\n<body>\n"
        html_document += f"<h1>{title}</h1>\n"
        html_document += f"<p><strong>Autor:</strong> {author}</p>\n"

        for term, data in self.glosario.items():
            html_document += f"<h2>{term}</h2>\n"
            html_document += f"<p><strong>Definición:</strong> {data['definicion']}</p>\n"
            html_document += "<p><strong>Ejemplos:</strong></p>\n<ul>\n"
            for example in data["ejemplos"]:
                html_document += f"<li>{example}</li>\n"
            html_document += "</ul>\n"

        html_document += "</body>\n</html>"

        return html_document

    def to_docx(self, author: str, title: str, path: str) -> str:
        """Convierte el glosario a un documento Word (docx)."""
        from docx import Document

        doc = Document()
        doc.add_heading(title, level=1)
        doc.add_paragraph(f"Autor: {author}")

        for term, data in self.glosario.items():
            doc.add_heading(term, level=2)
            doc.add_paragraph(f"Definición: {data['definicion']}")
            doc.add_paragraph("Ejemplos:")
            for example in data["ejemplos"]:
                doc.add_paragraph(f"- {example}", style="List Bullet")

        docx_filename = path + title + ".docx"
        doc.save(docx_filename)

        return docx_filename

    def to_xml(self, author: str, title: str) -> str:
        """Convierte el glosario a un documento XML."""
        xml_document = "<?xml version='1.0' encoding='UTF-8'?>\n"
        xml_document += f"<glosario>\n"
        xml_document += f"  <title>{title}</title>\n"
        xml_document += f"  <author>{author}</author>\n"

        for term, data in self.glosario.items():
            xml_document += f"  <term>\n"
            xml_document += f"    <name>{term}</name>\n"
            xml_document += f"    <definition>{data['definicion']}</definition>\n"
            xml_document += f"    <examples>\n"
            for example in data["ejemplos"]:
                xml_document += f"      <example>{example}</example>\n"
            xml_document += f"    </examples>\n"
            xml_document += f"  </term>\n"

        xml_document += f"</glosario>\n"

        return xml_document

    def from_xml(self, xml_path: str, encoding: str = "utf-8") -> dict:
        """Convierte un archivo XML a un diccionario de glosario."""
        import xml.etree.ElementTree as ET

        tree = ET.parse(xml_path)
        root = tree.getroot()

        self.glosario = {}

        for term in root.findall("term"):
            name = term.find("name").text
            definition = term.find("definition").text
            examples = [example.text for example in term.find("examples")]

            self.glosario[name] = {
                "definicion": definition,
                "ejemplos": examples
            }

        return self.glosario

# Ejemplo de uso
if __name__ == "__main__":
    glossary_manager = GlossaryManager()
    glossary_manager.load_glossary("input/glosario_agil.xml")

    to_latex = glossary_manager.to_latex("Moisés Leiva", "Glosario de Términos Ágiles")
    to_html = glossary_manager.to_html("Moisés Leiva", "Glosario de Términos Ágiles")
    to_md = glossary_manager.to_md("Moisés Leiva", "Glosario de Términos Ágiles")
    to_xml = glossary_manager.to_xml("Moisés Leiva", "Glosario de Términos Ágiles")


    # guardar latex, html, md y xml en /output/
    with open("output/glosario.tex", "w", encoding="utf-8") as file:
        file.write(to_latex)

    with open("output/glosario.html", "w", encoding="utf-8") as file:
        file.write(to_html)

    with open("output/glosario.md", "w", encoding="utf-8") as file:
        file.write(to_md)

    with open("output/glosario.xml", "w", encoding="utf-8") as file:
        file.write(to_xml)

    # Guardar en docx
    glossary_manager.to_docx("Moisés Leiva", "Glosario de Términos Ágiles", "output/")